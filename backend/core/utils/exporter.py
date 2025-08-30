import tempfile
from docx import Document
from reportlab.pdfgen import canvas

def export_docx(analysis):
    doc = Document()
    doc.add_heading("ATS Resume Analysis", 0)
    doc.add_paragraph(f"Score: {analysis.score}%")
    doc.add_paragraph("Suggestions:")
    for s in analysis.suggestions:
        doc.add_paragraph(f"- {s}")
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name

def export_pdf(analysis):
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    c = canvas.Canvas(tmp.name)
    c.drawString(100, 800, "ATS Resume Analysis")
    c.drawString(100, 780, f"Score: {analysis.score}%")
    y = 760
    for s in analysis.suggestions:
        c.drawString(100, y, f"- {s}")
        y -= 20
    c.save()
    return tmp.name
